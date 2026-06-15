"""Build a senior-partner crypto performance report.

The script reads the five-coin crypto panel, cleans the data, computes daily
returns with a zero risk-free rate, exports FT-style figures and tables, and
creates a Word report under ``fins2026/week2/scratch/crypto_report``.
"""

from __future__ import annotations

import os
import sys
import textwrap
from collections.abc import Iterable
from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd

MPLCONFIGDIR = Path(__file__).resolve().parent / "crypto_report" / ".matplotlib"
MPLCONFIGDIR.mkdir(parents=True, exist_ok=True)
os.environ.setdefault("MPLCONFIGDIR", str(MPLCONFIGDIR))

import matplotlib.dates as mdates  # noqa: E402
import matplotlib.pyplot as plt  # noqa: E402
import seaborn as sns  # noqa: E402
from matplotlib.ticker import FuncFormatter  # noqa: E402

DATA_URL = "https://openbondassetpricing.com/wp-content/uploads/2026/06/crypto_panel.csv"
TRADING_DAYS_PER_YEAR = 365
RISK_FREE_RATE = 0.0

SCRIPT_PATH = Path(__file__).resolve()
REPO_ROOT = SCRIPT_PATH.parents[3]
REPORT_ROOT = SCRIPT_PATH.parent / "crypto_report"
DATA_DIR = REPORT_ROOT / "data"
CODE_DIR = REPORT_ROOT / "code"
OUTPUT_DIR = REPORT_ROOT / "output"
FIGURE_DIR = OUTPUT_DIR / "figures"
TABLE_DIR = OUTPUT_DIR / "tables"
REPORT_DIR = REPORT_ROOT / "report"

RAW_CSV = DATA_DIR / "raw_crypto_panel.csv"
CLEAN_CSV = DATA_DIR / "crypto_panel_clean.csv"
CLEAN_PARQUET = DATA_DIR / "crypto_panel_clean.parquet"
RETURNS_CSV = DATA_DIR / "crypto_returns_wide.csv"
PRICES_CSV = DATA_DIR / "crypto_close_wide.csv"
REPORT_DOCX = REPORT_DIR / "crypto_investment_performance_report.docx"

if str(REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(REPO_ROOT))

from fintools.figures.export import FigureContext, export_figure_bundle  # noqa: E402
from fintools.figures.theme import FT_COLORS, figure_style  # noqa: E402


@dataclass(frozen=True)
class FigureAsset:
    """Metadata needed to insert an exported figure into the Word report."""

    path: Path
    context: FigureContext


def scaffold_report_folder() -> None:
    """Create a neat, self-contained analysis folder."""

    for folder in [DATA_DIR, CODE_DIR, FIGURE_DIR, TABLE_DIR, REPORT_DIR]:
        folder.mkdir(parents=True, exist_ok=True)

    readme = f"""# Crypto Report

This folder contains the Week 2 crypto investment performance report.

## Structure

- `data/` stores the cached raw panel, cleaned panel, close-price matrix, and return matrix.
- `code/` stores a copy of the executable analysis script used to build the report.
- `output/figures/` stores Word-ready FT-style PNG/PDF figures and caption sidecars.
- `output/tables/` stores final CSV tables rounded for reporting.
- `report/` stores the Word `.docx` report.

## Rebuild

From the repository root:

```bash
./.venv/bin/python fins2026/week2/scratch/crypto_narrative.py
```

The script assumes a zero risk-free rate and annualises daily crypto returns with
`{TRADING_DAYS_PER_YEAR}` calendar days per year.
"""
    (REPORT_ROOT / "README.md").write_text(readme, encoding="utf-8")

    agents = """# Crypto Report Agent Context

Scope: this context applies only inside `fins2026/week2/scratch/crypto_report`.

## Reporting Rules

- Keep the analysis self-contained in this folder; do not modify context files outside it.
- Treat `data/raw_crypto_panel.csv` as the cached source extract.
- Use `code/crypto_narrative.py` as the reproducible build script copy.
- Report returns, volatility, drawdowns, and correlations in percent terms where applicable.
- Round final table values to 3 decimal places.
- Keep table column names short enough for Word; avoid labels longer than two lines.
- Use a zero risk-free rate unless the senior partner explicitly changes the mandate.
- Keep figure notes short; do not let text below figures overrun the image.
- Prefer FT-style charts with direct, decision-useful titles.

## Rebuild Command

```bash
./.venv/bin/python fins2026/week2/scratch/crypto_narrative.py
```
"""
    (REPORT_ROOT / "AGENTS.md").write_text(agents, encoding="utf-8")


def load_crypto_panel() -> pd.DataFrame:
    """Load the raw panel from the local cache or remote URL."""

    if RAW_CSV.exists():
        raw = pd.read_csv(RAW_CSV)
    else:
        raw = pd.read_csv(DATA_URL)
        raw.to_csv(RAW_CSV, index=False)
    return raw


def clean_crypto_panel(raw: pd.DataFrame) -> pd.DataFrame:
    """Clean the crypto panel and compute simple daily returns."""

    panel = raw.copy()
    panel.columns = [str(col).strip().lower() for col in panel.columns]

    required = {"date", "ticker", "close"}
    missing = sorted(required.difference(panel.columns))
    if missing:
        raise ValueError(f"Missing required columns: {missing}")

    panel["date"] = pd.to_datetime(panel["date"], errors="coerce")
    panel["ticker"] = panel["ticker"].astype(str).str.strip().str.upper()

    numeric_columns = [col for col in panel.columns if col not in {"date", "ticker"}]
    panel[numeric_columns] = panel[numeric_columns].apply(pd.to_numeric, errors="coerce")

    panel = panel.dropna(subset=["date", "ticker", "close"])
    panel = panel.sort_values(["ticker", "date"]).reset_index(drop=True)
    before = len(panel)
    panel = panel.drop_duplicates(subset=["date", "ticker"], keep="last").reset_index(drop=True)
    dropped = before - len(panel)
    if dropped:
        print(f"Dropped duplicate date-ticker rows: {dropped:,}")

    panel["daily_return"] = panel.groupby("ticker", observed=True)["close"].pct_change()
    if "market_cap" in panel.columns:
        panel["cap_weight_proxy"] = panel["market_cap"] / panel.groupby("date")[
            "market_cap"
        ].transform("sum")
    if "share_volume" in panel.columns:
        panel["dollar_volume"] = panel["share_volume"] * panel["close"]
    elif "volume" in panel.columns:
        panel["dollar_volume"] = panel["volume"] * panel["close"]

    return panel


def build_wide_tables(panel: pd.DataFrame) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Return wide close prices and daily returns."""

    close_wide = panel.pivot_table(index="date", columns="ticker", values="close").sort_index()
    returns_wide = close_wide.pct_change()
    returns_wide = returns_wide.dropna(how="all")
    return close_wide, returns_wide


def max_drawdown(series: pd.Series) -> float:
    """Return max drawdown for a return series in decimal units."""

    wealth = (1.0 + series.dropna()).cumprod()
    if wealth.empty:
        return np.nan
    return float((wealth / wealth.cummax() - 1.0).min())


def build_summary_tables(
    panel: pd.DataFrame,
    close_wide: pd.DataFrame,
    returns_wide: pd.DataFrame,
) -> dict[str, pd.DataFrame]:
    """Build concise final tables for the report."""

    observations = returns_wide.count()
    avg_daily = returns_wide.mean()
    ann_return = avg_daily * TRADING_DAYS_PER_YEAR
    daily_vol = returns_wide.std()
    ann_vol = daily_vol * np.sqrt(TRADING_DAYS_PER_YEAR)
    sharpe = np.where(ann_vol > 0, ann_return / ann_vol, np.nan)
    cumulative = (1.0 + returns_wide).prod(skipna=True) - 1.0
    drawdown = returns_wide.apply(max_drawdown)
    hit_rate = (returns_wide > RISK_FREE_RATE).sum() / observations
    start_price = close_wide.apply(
        lambda col: col.dropna().iloc[0] if col.notna().any() else np.nan
    )
    end_price = close_wide.apply(lambda col: col.dropna().iloc[-1] if col.notna().any() else np.nan)

    perf = pd.DataFrame(
        {
            "Obs": observations,
            "Avg %": avg_daily * 100,
            "Ann %": ann_return * 100,
            "Vol %": ann_vol * 100,
            "Sharpe": sharpe,
            "Cum %": cumulative * 100,
            "MDD %": drawdown * 100,
            "Up %": hit_rate * 100,
            "Start": start_price,
            "End": end_price,
        }
    ).round(3)
    perf.index.name = "Coin"

    corr = returns_wide.corr().round(3)
    corr.index.name = "Coin"

    monthly = (1.0 + returns_wide).resample("ME").prod() - 1.0
    monthly_stats = pd.DataFrame(
        {
            "Avg %": monthly.mean() * 100,
            "Vol %": monthly.std() * 100,
            "Best %": monthly.max() * 100,
            "Worst %": monthly.min() * 100,
        }
    ).round(3)
    monthly_stats.index.name = "Coin"

    if "dollar_volume" in panel.columns:
        liquidity = (
            panel.groupby("ticker", observed=True)["dollar_volume"]
            .agg(["mean", "median"])
            .rename(columns={"mean": "Avg $vol", "median": "Med $vol"})
            / 1_000_000
        ).round(3)
        liquidity.index.name = "Coin"
    else:
        liquidity = pd.DataFrame(index=perf.index)

    return {
        "performance": perf,
        "correlation": corr,
        "monthly": monthly_stats,
        "liquidity": liquidity,
    }


def save_tables(tables: dict[str, pd.DataFrame]) -> None:
    """Write final report tables to CSV."""

    for name, table in tables.items():
        table.to_csv(TABLE_DIR / f"{name}.csv")


def sample_label(dates: Iterable[pd.Timestamp]) -> str:
    """Format the sample period for captions."""

    values = pd.Series(list(dates)).dropna()
    if values.empty:
        return ""
    return f"{values.min().date()} to {values.max().date()}"


def wrap_note(text: str, width: int = 118) -> str:
    """Keep figure notes short and within the figure canvas."""

    return "\n".join(textwrap.wrap(text, width=width))


def add_source_note(fig: plt.Figure, text: str) -> None:
    """Add a compact source note below a figure."""

    fig.text(
        0.01,
        0.01,
        wrap_note(text),
        ha="left",
        va="bottom",
        fontsize=8,
        color=FT_COLORS["muted"],
    )


def export_current_figure(fig: plt.Figure, stem: str, context: FigureContext) -> FigureAsset:
    """Export one figure as PNG/PDF and return report metadata."""

    paths = export_figure_bundle(fig, FIGURE_DIR, stem, context=context, formats=("png", "pdf"))
    plt.close(fig)
    return FigureAsset(path=paths["png"], context=context)


def dollar_formatter(value: float, _position: int) -> str:
    """Format growth-of-one-dollar axes."""

    if value >= 10:
        return f"${value:,.0f}"
    return f"${value:,.1f}"


def percent_formatter(value: float, _position: int) -> str:
    """Format decimal axis values as percent."""

    return f"{value * 100:.0f}%"


def make_figures(
    close_wide: pd.DataFrame,
    returns_wide: pd.DataFrame,
    tables: dict[str, pd.DataFrame],
) -> list[FigureAsset]:
    """Create the FT-style figure pack."""

    figures: list[FigureAsset] = []
    sample = sample_label(close_wide.index)
    colors = FT_COLORS

    with figure_style("word_a4", style="ft", ft_background=True):
        wealth = (1.0 + returns_wide.fillna(0.0)).cumprod()
        fig, ax = plt.subplots(figsize=(6.27, 3.75))
        for coin in wealth.columns:
            ax.plot(wealth.index, wealth[coin], label=coin, linewidth=1.7, alpha=0.92)
        ax.set_title("Crypto growth of $1")
        ax.set_xlabel("")
        ax.set_ylabel("Portfolio value")
        ax.set_yscale("log")
        ax.yaxis.set_major_formatter(FuncFormatter(dollar_formatter))
        ax.grid(axis="y")
        ax.legend(ncol=min(5, len(wealth.columns)), loc="upper center", bbox_to_anchor=(0.5, -0.12))
        ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=4, maxticks=7))
        add_source_note(fig, f"Source: OpenBond Asset Pricing crypto panel. Sample: {sample}.")
        context = FigureContext(
            title="Growth of one dollar by coin",
            note="Compounds daily simple returns for each coin.",
            sample=sample,
            units="Index, first valid observation equals $1; log scale.",
            source="OpenBond Asset Pricing crypto panel.",
        )
        figures.append(export_current_figure(fig, "fig01_growth_of_one_dollar", context))

    with figure_style("word_a4", style="ft", ft_background=True):
        drawdowns = (1.0 + returns_wide.fillna(0.0)).cumprod()
        drawdowns = drawdowns / drawdowns.cummax() - 1.0
        fig, ax = plt.subplots(figsize=(6.27, 3.75))
        for coin in drawdowns.columns:
            ax.plot(drawdowns.index, drawdowns[coin], label=coin, linewidth=1.4, alpha=0.82)
        ax.set_title("Drawdowns show repeated deep capital impairment")
        ax.set_xlabel("")
        ax.set_ylabel("Drawdown")
        ax.yaxis.set_major_formatter(FuncFormatter(percent_formatter))
        ax.grid(axis="y")
        ax.legend(
            ncol=min(5, len(drawdowns.columns)),
            loc="upper center",
            bbox_to_anchor=(0.5, -0.12),
        )
        ax.xaxis.set_major_locator(mdates.AutoDateLocator(minticks=4, maxticks=7))
        add_source_note(fig, f"Source: OpenBond Asset Pricing crypto panel. Sample: {sample}.")
        context = FigureContext(
            title="Drawdown paths by coin",
            note="Drawdown is the percentage fall from each coin's previous wealth peak.",
            sample=sample,
            units="Percent.",
            source="OpenBond Asset Pricing crypto panel.",
        )
        figures.append(export_current_figure(fig, "fig02_drawdowns", context))

    with figure_style("word_a4", style="ft", ft_background=True):
        ranking = tables["performance"].sort_values("Sharpe", ascending=True)
        fig, ax = plt.subplots(figsize=(6.27, 3.75))
        y = np.arange(len(ranking))
        ax.hlines(y=y, xmin=0, xmax=ranking["Sharpe"], color=colors["axis"], linewidth=2)
        ax.scatter(ranking["Sharpe"], y, s=70, color=colors["maroon"], zorder=3)
        for ypos, value in zip(y, ranking["Sharpe"], strict=False):
            ax.text(value, ypos, f" {value:.2f}", va="center", fontsize=9)
        ax.set_yticks(y)
        ax.set_yticklabels(ranking.index)
        ax.axvline(0, color=colors["charcoal"], linewidth=0.8)
        ax.set_title("Sharpe ratios are uneven across the five coins")
        ax.set_xlabel("Annualised Sharpe ratio, rf = 0")
        ax.grid(axis="x")
        add_source_note(fig, f"Source: OpenBond Asset Pricing crypto panel. Sample: {sample}.")
        context = FigureContext(
            title="Risk-adjusted performance ranking",
            note="Annualised Sharpe ratio uses a zero risk-free rate.",
            sample=sample,
            units="Ratio.",
            source="OpenBond Asset Pricing crypto panel.",
        )
        figures.append(export_current_figure(fig, "fig03_sharpe_ranking", context))

    with figure_style("word_a4", style="ft", ft_background=True):
        risk_return = tables["performance"]
        fig, ax = plt.subplots(figsize=(6.27, 3.75))
        ax.scatter(
            risk_return["Vol %"],
            risk_return["Ann %"],
            s=90,
            color=colors["blue"],
            alpha=0.9,
        )
        for coin, row in risk_return.iterrows():
            ax.annotate(
                coin,
                (row["Vol %"], row["Ann %"]),
                xytext=(5, 5),
                textcoords="offset points",
                fontsize=9,
            )
        ax.axhline(0, color=colors["axis"], linewidth=0.8)
        ax.set_title("Return compensation versus annualised volatility")
        ax.set_xlabel("Vol %")
        ax.set_ylabel("Ann %")
        ax.grid(True)
        add_source_note(fig, f"Source: OpenBond Asset Pricing crypto panel. Sample: {sample}.")
        context = FigureContext(
            title="Annualised return and volatility by coin",
            note="Return and volatility are annualised from daily simple returns.",
            sample=sample,
            units="Percent.",
            source="OpenBond Asset Pricing crypto panel.",
        )
        figures.append(export_current_figure(fig, "fig04_risk_return_map", context))

    with figure_style("word_a4", style="ft", ft_background=True):
        corr = returns_wide.corr()
        fig, ax = plt.subplots(figsize=(6.27, 4.2))
        sns.heatmap(
            corr,
            ax=ax,
            vmin=-1,
            vmax=1,
            cmap=sns.diverging_palette(240, 10, as_cmap=True),
            annot=True,
            fmt=".2f",
            square=True,
            linewidths=0.8,
            cbar_kws={"label": "Correlation"},
        )
        ax.set_title("Daily return correlations are high but not identical")
        ax.set_xlabel("")
        ax.set_ylabel("")
        add_source_note(fig, f"Source: OpenBond Asset Pricing crypto panel. Sample: {sample}.")
        context = FigureContext(
            title="Daily return correlation matrix",
            note="Correlations use pairwise complete daily returns.",
            sample=sample,
            units="Correlation coefficient.",
            source="OpenBond Asset Pricing crypto panel.",
        )
        figures.append(export_current_figure(fig, "fig05_correlation_heatmap", context))

    with figure_style("word_a4", style="ft", ft_background=True):
        monthly = (1.0 + returns_wide).resample("ME").prod() - 1.0
        monthly_long = monthly.stack().rename("monthly_return").reset_index()
        monthly_long.columns = ["date", "coin", "monthly_return"]
        fig, ax = plt.subplots(figsize=(6.27, 3.75))
        sns.boxplot(
            data=monthly_long,
            x="coin",
            y="monthly_return",
            ax=ax,
            color=colors["light_blue"],
            fliersize=2.5,
            linewidth=0.9,
        )
        ax.axhline(0, color=colors["charcoal"], linewidth=0.8)
        ax.set_title("Monthly return distributions retain large downside tails")
        ax.set_xlabel("")
        ax.set_ylabel("Monthly return")
        ax.yaxis.set_major_formatter(FuncFormatter(percent_formatter))
        ax.grid(axis="y")
        add_source_note(fig, f"Source: OpenBond Asset Pricing crypto panel. Sample: {sample}.")
        context = FigureContext(
            title="Monthly return distribution by coin",
            note="Boxes show the interquartile range of compounded monthly returns.",
            sample=sample,
            units="Percent.",
            source="OpenBond Asset Pricing crypto panel.",
        )
        figures.append(export_current_figure(fig, "fig06_monthly_return_distribution", context))

    return figures


def format_table_for_word(table: pd.DataFrame) -> pd.DataFrame:
    """Convert a dataframe into display-safe strings for Word."""

    display = table.reset_index().copy()
    for col in display.columns:
        if pd.api.types.is_numeric_dtype(display[col]):
            display[col] = display[col].map(lambda value: "" if pd.isna(value) else f"{value:,.3f}")
    return display


def add_word_table(document: object, table: pd.DataFrame) -> None:
    """Add a compact Word table with short headers."""

    display = format_table_for_word(table)
    word_table = document.add_table(rows=1, cols=len(display.columns))
    word_table.style = "Table Grid"
    for cell, column in zip(word_table.rows[0].cells, display.columns, strict=False):
        cell.text = str(column)
    for _, row in display.iterrows():
        cells = word_table.add_row().cells
        for cell, value in zip(cells, row, strict=False):
            cell.text = str(value)


def add_figure(document: object, figure: FigureAsset) -> None:
    """Insert an inline figure and a short caption into Word."""

    from docx.shared import Inches, Pt

    document.add_picture(str(figure.path), width=Inches(6.25))
    caption = document.add_paragraph(style="Caption")
    caption.paragraph_format.space_after = Pt(6)
    caption.add_run(figure.context.caption_text())


def write_report(
    panel: pd.DataFrame,
    close_wide: pd.DataFrame,
    tables: dict[str, pd.DataFrame],
    figures: list[FigureAsset],
) -> None:
    """Generate the Word report."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    sample = sample_label(close_wide.index)
    perf = tables["performance"].sort_values("Sharpe", ascending=False)
    best = perf.index[0]
    worst = perf.index[-1]
    best_sharpe = perf.iloc[0]["Sharpe"]
    worst_drawdown_coin = tables["performance"]["MDD %"].idxmin()
    worst_drawdown = tables["performance"].loc[worst_drawdown_coin, "MDD %"]
    mean_corr = (
        tables["correlation"]
        .where(~np.eye(len(tables["correlation"]), dtype=bool))
        .stack()
        .mean()
    )

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 12)]:
        styles[style_name].font.name = "Aptos"
        styles[style_name].font.size = Pt(size)
        styles[style_name].font.bold = True

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Crypto Investment Performance: Five-Coin Panel")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Senior partner briefing | Sample: {sample} | Risk-free rate: 0")

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        f"{best} has the strongest risk-adjusted performance in the sample, with an "
        f"annualised Sharpe ratio of {best_sharpe:.3f}; {worst} sits at the bottom "
        "of the same ranking. The result is not a low-risk story: annualised "
        "volatility is high for every coin, and the maximum drawdown reaches "
        f"{worst_drawdown:.3f}% for {worst_drawdown_coin}."
    )
    document.add_paragraph(
        f"Diversification within the five-coin set is limited. The average off-diagonal "
        f"daily return correlation is {mean_corr:.3f}, so coin selection changes relative "
        "winners but does not remove the common crypto risk cycle."
    )

    document.add_heading("Data and Method", level=1)
    document.add_paragraph(
        f"The analysis uses the OpenBond Asset Pricing crypto panel from {sample}. "
        f"The cleaned panel contains {len(panel):,} date-coin observations and "
        f"{close_wide.shape[1]} coins. Daily simple returns are computed from close "
        "prices by coin. All excess-return and Sharpe calculations assume a zero "
        "risk-free rate. Annualised returns use the mean daily return multiplied by "
        f"{TRADING_DAYS_PER_YEAR}; annualised volatility scales daily volatility by "
        f"the square root of {TRADING_DAYS_PER_YEAR}."
    )

    document.add_heading("Performance Scorecard", level=1)
    document.add_paragraph(
        "Table 1 reports the core investment scorecard. Return, volatility, cumulative "
        "return, maximum drawdown, and up-day frequency are shown in percent; the Sharpe "
        "ratio is unitless."
    )
    add_word_table(document, tables["performance"])
    document.add_paragraph("Table 1. Five-coin performance scorecard.", style="Caption")

    document.add_heading("Price and Return Dynamics", level=1)
    for figure in figures[:4]:
        add_figure(document, figure)
    document.add_paragraph(
        "The growth and drawdown charts show why the ranking must be read jointly with "
        "path risk. High compounded gains can coexist with deep interim losses, so the "
        "portfolio decision should treat drawdown tolerance as a binding constraint."
    )

    document.add_heading("Correlation and Distribution Risk", level=1)
    for figure in figures[4:]:
        add_figure(document, figure)
    document.add_paragraph(
        "The correlation matrix shows a shared risk factor across coins. The monthly "
        "distribution view confirms that tail risk remains visible even after daily "
        "returns are compounded to a lower frequency."
    )

    document.add_heading("Correlation Matrix", level=1)
    add_word_table(document, tables["correlation"])
    document.add_paragraph("Table 2. Daily return correlations.", style="Caption")

    document.add_heading("Monthly Return Diagnostics", level=1)
    add_word_table(document, tables["monthly"])
    document.add_paragraph("Table 3. Compounded monthly return diagnostics.", style="Caption")

    if not tables["liquidity"].empty:
        document.add_heading("Liquidity Proxy", level=1)
        document.add_paragraph(
            "Table 4 reports average and median dollar volume in millions where the "
            "source panel includes a volume field."
        )
        add_word_table(document, tables["liquidity"])
        document.add_paragraph("Table 4. Dollar volume proxy in millions.", style="Caption")

    document.add_heading("Conclusion", level=1)
    document.add_paragraph(
        f"The five-coin panel rewards selectivity but not complacency. {best} leads on "
        "risk-adjusted performance, while the common correlation structure and deep "
        "drawdowns mean a crypto allocation should be sized against portfolio-level "
        "loss capacity rather than headline cumulative return alone."
    )

    document.core_properties.title = "Crypto Investment Performance: Five-Coin Panel"
    document.core_properties.subject = "FINS2026 crypto performance report"
    document.core_properties.author = "FINS2026 analysis pipeline"
    document.save(REPORT_DOCX)


def main() -> int:
    """Run the full crypto report build."""

    scaffold_report_folder()
    raw = load_crypto_panel()
    panel = clean_crypto_panel(raw)
    close_wide, returns_wide = build_wide_tables(panel)
    tables = build_summary_tables(panel, close_wide, returns_wide)

    panel.to_csv(CLEAN_CSV, index=False)
    panel.to_parquet(CLEAN_PARQUET, index=False)
    close_wide.to_csv(PRICES_CSV)
    returns_wide.to_csv(RETURNS_CSV)
    save_tables(tables)

    figures = make_figures(close_wide, returns_wide, tables)
    write_report(panel, close_wide, tables, figures)

    CODE_DIR.mkdir(parents=True, exist_ok=True)
    (CODE_DIR / "crypto_narrative.py").write_text(
        SCRIPT_PATH.read_text(encoding="utf-8"),
        encoding="utf-8",
    )

    print("Crypto report build complete")
    print(f"Clean data: {CLEAN_CSV}")
    print(f"Figures: {FIGURE_DIR}")
    print(f"Tables: {TABLE_DIR}")
    print(f"Word report: {REPORT_DOCX}")
    print("\nPerformance table")
    print(tables["performance"].to_string())
    print("\nCorrelation table")
    print(tables["correlation"].to_string())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
